import streamlit as st
from utils.gpt_api import get_api_response, prompt_original
from docx import Document
from io import BytesIO
from datetime import datetime

st.set_page_config(
    page_title="Corregidor by CMR",
    page_icon="🚩"
)

st.markdown("""<h1 style="text-align: center; color: black;">
    Corregidor<sub>by CMR</sub>
</h1>""", unsafe_allow_html=True)
st.markdown('---')

## prompt modificable - input texto
with st.form('form_1'):
    images = st.file_uploader('Suba todas las imágenes que desee analizar', type= ['.jpg', '.jpeg'], accept_multiple_files=True)
    prompt = st.text_area('Introduzca el prompt que desea ejecutar:')
    ok_button = st.form_submit_button('Ok')

if ok_button:
    ## ejecutar prompt
    my_bar = st.progress(0, text='Leyendo imágenes y generando respuestas...')

    # images = os.listdir('./data')
    images = [ (image.read(), image.name) for image in images if '.jpg' in image.name or '.jpeg'in image.name]

    # initialize doc to save as docx
    doc = Document()

    divide_percent = len(images)
    for i in range(divide_percent):
        image, image_name = images[i]

        with st.spinner('Esperando la respuesta del modelo...'):
            response = get_api_response(prompt, image) 
        
        with st.spinner('Guardando la respuesta...'):
            response_md = response.json()['choices'][0]['message']['content']
            doc.add_heading(f'Imagen {i+1}')
            doc.add_paragraph(response_md)
            if i+1 != divide_percent:
                doc.add_page_break()

        my_bar.progress((i+1)/divide_percent, text='Procesando ficheros')



    ## resultados análisis (nr ficheros procesados)
    st.success(f'{len(images)} ficheros procesados correctamente')


    # Guardar el documento en memoria
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    time_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    st.download_button(
        label="📥 Descargar resumen",
        data=doc_io,
        file_name=f"resultado_corregidor_{time_stamp}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



